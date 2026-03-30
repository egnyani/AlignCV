"""Strict validators for locked-template resume generation."""

from __future__ import annotations

import math
import zipfile
from collections import Counter
from pathlib import Path
from typing import Any

from docx import Document
from lxml import etree

from pipeline.pdf_converter import pdf_page_count


_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def _nonempty_paragraphs(document: Document) -> list[Any]:
    return [p for p in document.paragraphs if (p.text or "").strip()]


def _style_counts(document: Document) -> dict[str, int]:
    counts = Counter()
    for paragraph in _nonempty_paragraphs(document):
        style = paragraph.style.name if paragraph.style else "UNKNOWN"
        counts[style] += 1
    return dict(counts)


_STYLE_LINE_CAPACITY = {
    "Title": 28,
    "Body Text": 108,
    "Heading 1": 40,
    "Normal": 100,
    "List Paragraph": 92,
}

_STYLE_LINE_HEIGHT_TWIPS = {
    "Title": 360,
    "Body Text": 240,
    "Heading 1": 280,
    "Normal": 240,
    "List Paragraph": 240,
}

_TWIP_TO_EMU = 635


def _section_order(document: Document) -> list[str]:
    order: list[str] = []
    for paragraph in _nonempty_paragraphs(document):
        text = (paragraph.text or "").strip()
        style = paragraph.style.name if paragraph.style else ""
        if style.startswith("Heading") or (text.isupper() and len(text.split()) <= 4):
            order.append(text)
    return order


def _nonempty_texts(document: Document) -> list[str]:
    return [(p.text or "").strip() for p in _nonempty_paragraphs(document)]


def _margins(document: Document) -> list[tuple[int, int, int, int]]:
    values: list[tuple[int, int, int, int]] = []
    for section in document.sections:
        values.append((
            int(section.top_margin or 0),
            int(section.bottom_margin or 0),
            int(section.left_margin or 0),
            int(section.right_margin or 0),
        ))
    return values


def _page_sizes(document: Document) -> list[tuple[int, int]]:
    values: list[tuple[int, int]] = []
    for section in document.sections:
        values.append((int(section.page_width or 0), int(section.page_height or 0)))
    return values


def _xml_root(docx_path: str | Path) -> etree._Element:
    with zipfile.ZipFile(str(docx_path), "r") as zf:
        return etree.fromstring(zf.read("word/document.xml"))


def _sect_pr_fingerprints(docx_path: str | Path) -> list[dict[str, str]]:
    root = _xml_root(docx_path)
    results: list[dict[str, str]] = []
    for sect in root.xpath(".//w:sectPr", namespaces=_NS):
        item: dict[str, str] = {}
        for key, tag in {
            "page_size": "w:pgSz",
            "page_margins": "w:pgMar",
            "page_borders": "w:pgBorders",
            "columns": "w:cols",
        }.items():
            node = sect.find(tag, namespaces=_NS)
            item[key] = etree.tostring(node, encoding="unicode") if node is not None else ""
        results.append(item)
    return results


def _paragraph_fingerprints(docx_path: str | Path) -> list[dict[str, Any]]:
    root = _xml_root(docx_path)
    fingerprints: list[dict[str, Any]] = []
    for para in root.xpath(".//w:body/w:p", namespaces=_NS):
        text = "".join(para.xpath(".//w:t/text()", namespaces=_NS)).strip()
        if not text:
            continue
        ppr = para.find("w:pPr", namespaces=_NS)
        fingerprints.append(
            {
                "text": text,
                "tabs": text.count("\t"),
                "ppr": etree.tostring(ppr, encoding="unicode") if ppr is not None else "",
                "run_count": len(para.xpath("./w:r", namespaces=_NS)),
            }
        )
    return fingerprints


def _contact_structure(document: Document, docx_path: str | Path) -> dict[str, Any]:
    paragraphs = _nonempty_paragraphs(document)
    xml_paragraphs = _paragraph_fingerprints(docx_path)
    contact = paragraphs[1] if len(paragraphs) > 1 else None
    xml_contact = xml_paragraphs[1] if len(xml_paragraphs) > 1 else {}
    return {
        "name_text": (paragraphs[0].text if paragraphs else "").strip(),
        "contact_text": (contact.text if contact else "").strip(),
        "name_style": paragraphs[0].style.name if paragraphs and paragraphs[0].style else "",
        "contact_style": contact.style.name if contact and contact.style else "",
        "contact_tabs": xml_contact.get("tabs", 0),
        "contact_ppr": xml_contact.get("ppr", ""),
    }


def _estimated_line_usage(document: Document) -> int:
    total = 0
    for paragraph in _nonempty_paragraphs(document):
        style = paragraph.style.name if paragraph.style else "Normal"
        capacity = _STYLE_LINE_CAPACITY.get(style, 95)
        total += max(1, math.ceil(len((paragraph.text or "").strip()) / capacity))
    return total


def _paragraph_line_counts(document: Document) -> list[int]:
    counts: list[int] = []
    for paragraph in _nonempty_paragraphs(document):
        style = paragraph.style.name if paragraph.style else "Normal"
        capacity = _STYLE_LINE_CAPACITY.get(style, 95)
        counts.append(max(1, math.ceil(len((paragraph.text or "").strip()) / capacity)))
    return counts


def _usable_page_height(document: Document) -> int:
    if not document.sections:
        return 0
    section = document.sections[0]
    return int(section.page_height or 0) - int(section.top_margin or 0) - int(section.bottom_margin or 0)


def _paragraph_estimated_height_emu(paragraph: Any) -> int:
    style = paragraph.style.name if paragraph.style else "Normal"
    capacity = _STYLE_LINE_CAPACITY.get(style, 95)
    line_height = _STYLE_LINE_HEIGHT_TWIPS.get(style, 240)
    text = (paragraph.text or "").strip()
    line_count = max(1, math.ceil(len(text) / capacity))
    spacing_before = int(paragraph.paragraph_format.space_before or 0)
    spacing_after = int(paragraph.paragraph_format.space_after or 0)
    line_height_emu = line_height * _TWIP_TO_EMU
    return (line_count * line_height_emu) + spacing_before + spacing_after


def _estimated_content_height_emu(document: Document) -> int:
    return sum(_paragraph_estimated_height_emu(p) for p in _nonempty_paragraphs(document))


def validate_layout_preservation(
    template_path: str | Path,
    output_path: str | Path,
    modified_sections: list[str] | None = None,
    pdf_path: str | Path | None = None,
) -> dict[str, Any]:
    """Compare output DOCX structure against the locked template."""
    template_doc = Document(str(template_path))
    output_doc = Document(str(output_path))

    template_paragraphs = _nonempty_paragraphs(template_doc)
    output_paragraphs = _nonempty_paragraphs(output_doc)
    template_texts = _nonempty_texts(template_doc)
    output_texts = _nonempty_texts(output_doc)
    template_styles = _style_counts(template_doc)
    output_styles = _style_counts(output_doc)
    template_order = _section_order(template_doc)
    output_order = _section_order(output_doc)
    template_margins = _margins(template_doc)
    output_margins = _margins(output_doc)
    template_page_sizes = _page_sizes(template_doc)
    output_page_sizes = _page_sizes(output_doc)
    template_sect = _sect_pr_fingerprints(template_path)
    output_sect = _sect_pr_fingerprints(output_path)
    template_contact = _contact_structure(template_doc, template_path)
    output_contact = _contact_structure(output_doc, output_path)
    template_line_usage = _estimated_line_usage(template_doc)
    output_line_usage = _estimated_line_usage(output_doc)
    template_line_counts = _paragraph_line_counts(template_doc)
    output_line_counts = _paragraph_line_counts(output_doc)
    usable_height = _usable_page_height(template_doc)
    template_content_height = _estimated_content_height_emu(template_doc)
    output_content_height = _estimated_content_height_emu(output_doc)

    template_summary_heading_count = sum(1 for text in template_texts if text == "SUMMARY")
    output_summary_heading_count = sum(1 for text in output_texts if text == "SUMMARY")
    template_summary_heading_index = template_texts.index("SUMMARY") if "SUMMARY" in template_texts else -1
    output_summary_heading_index = output_texts.index("SUMMARY") if "SUMMARY" in output_texts else -1
    template_pre_summary = template_texts[:template_summary_heading_index]
    output_pre_summary = output_texts[:output_summary_heading_index]
    template_summary_body_index = template_summary_heading_index + 1 if template_summary_heading_index >= 0 else -1
    output_summary_body_index = output_summary_heading_index + 1 if output_summary_heading_index >= 0 else -1
    template_summary_body = template_texts[template_summary_body_index] if template_summary_body_index < len(template_texts) else ""
    output_summary_body = output_texts[output_summary_body_index] if output_summary_body_index < len(output_texts) else ""

    style_deltas: dict[str, dict[str, int]] = {}
    for style in sorted(set(template_styles) | set(output_styles)):
        before = template_styles.get(style, 0)
        after = output_styles.get(style, 0)
        if before != after:
            style_deltas[style] = {"template": before, "output": after}

    borders_same = template_sect == output_sect
    contact_structure_same = template_contact == output_contact
    page_count = pdf_page_count(Path(pdf_path)) if pdf_path else None
    content_fill_ratio = round(output_content_height / max(template_content_height, 1), 3)
    bottom_whitespace_emu = max(0, template_content_height - output_content_height)
    template_fill_ratio = round(template_content_height / max(usable_height, 1), 3)
    # Treat the original template as the target fill profile and allow a modest
    # band around it so we don't leave a large empty footer or overflow badly.
    bottom_whitespace_reasonable = 0.9 <= content_fill_ratio <= 1.03
    content_within_template_bounds = output_content_height <= int(template_content_height * 1.03)

    return {
        "template": str(Path(template_path)),
        "output": str(Path(output_path)),
        "modified_sections": modified_sections or [],
        "page_size_same": template_page_sizes == output_page_sizes,
        "template_page_sizes": template_page_sizes,
        "output_page_sizes": output_page_sizes,
        "margins_same": template_margins == output_margins,
        "template_margins": template_margins,
        "output_margins": output_margins,
        "section_count_same": len(template_doc.sections) == len(output_doc.sections),
        "template_section_count": len(template_doc.sections),
        "output_section_count": len(output_doc.sections),
        "borders_same": borders_same,
        "paragraph_count_same": len(template_paragraphs) == len(output_paragraphs),
        "template_paragraph_count": len(template_paragraphs),
        "output_paragraph_count": len(output_paragraphs),
        "major_styles_same": not style_deltas,
        "style_deltas": style_deltas,
        "section_order_same": template_order == output_order,
        "template_section_order": template_order,
        "output_section_order": output_order,
        "contact_header_structure_same": contact_structure_same,
        "template_contact_header": template_contact,
        "output_contact_header": output_contact,
        "template_estimated_line_usage": template_line_usage,
        "output_estimated_line_usage": output_line_usage,
        "line_counts_same": template_line_counts == output_line_counts,
        "template_paragraph_line_counts": template_line_counts,
        "output_paragraph_line_counts": output_line_counts,
        "usable_page_height_emu": usable_height,
        "template_content_height_emu": template_content_height,
        "output_content_height_emu": output_content_height,
        "bottom_whitespace_emu": bottom_whitespace_emu,
        "template_fill_ratio": template_fill_ratio,
        "content_fill_ratio": content_fill_ratio,
        "bottom_whitespace_reasonable": bottom_whitespace_reasonable,
        "content_within_template_bounds": content_within_template_bounds,
        "summary_heading_count_same": template_summary_heading_count == output_summary_heading_count == 1,
        "template_summary_heading_count": template_summary_heading_count,
        "output_summary_heading_count": output_summary_heading_count,
        "summary_heading_position_same": template_summary_heading_index == output_summary_heading_index,
        "template_summary_heading_index": template_summary_heading_index,
        "output_summary_heading_index": output_summary_heading_index,
        "top_paragraphs_before_summary_same": template_pre_summary == output_pre_summary,
        "template_pre_summary_paragraphs": template_pre_summary,
        "output_pre_summary_paragraphs": output_pre_summary,
        "summary_body_in_place": template_summary_body_index == output_summary_body_index and bool(output_summary_body),
        "template_summary_body_index": template_summary_body_index,
        "output_summary_body_index": output_summary_body_index,
        "template_summary_body": template_summary_body,
        "output_summary_body": output_summary_body,
        "one_page": page_count == 1 if page_count is not None else None,
        "page_count": page_count,
    }


def layout_validation_passed(validation: dict[str, Any]) -> bool:
    required = [
        "page_size_same",
        "margins_same",
        "section_count_same",
        "borders_same",
        "paragraph_count_same",
        "major_styles_same",
        "section_order_same",
        "contact_header_structure_same",
        "summary_heading_count_same",
        "summary_heading_position_same",
        "top_paragraphs_before_summary_same",
        "summary_body_in_place",
        "line_counts_same",
        "content_within_template_bounds",
        "bottom_whitespace_reasonable",
        "one_page",
    ]
    return all(validation.get(key) is True for key in required)
